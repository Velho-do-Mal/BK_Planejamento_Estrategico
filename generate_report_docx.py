#!/usr/bin/env python3
"""
generate_report_docx.py

Gera um relatório Word (relatorio_planejamento.docx) a partir do arquivo planning.json
(se existir na pasta) ou de dados de exemplo.

Dependências:
  pip install python-docx pandas plotly kaleido numpy

Uso:
  python generate_report_docx.py
"""

import json
import io
import os
from datetime import datetime
from typing import Dict, Any, List

import numpy as np
import pandas as pd
import plotly.graph_objs as go
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# -------------------------
# Utilitários para carregar/normalizar dados
# -------------------------

def load_planning_json(path: str = "planning.json") -> Dict[str, Any]:
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data
    # Exemplo mínimo se não existir
    example = {
        "partners": [
            {"nome":"Márcio Nunes Knopp","cargo":"Sócio-Diretor","email":"marcio@bk-engenharia.com","telefone":"+55 (00) 99999-9999","observacoes":""}
        ],
        "areas": [
            {"area":"Projetos Elétricos","responsavel":"Márcio","email":"marcio@bk-engenharia.com","observacoes":""}
        ],
        "swot": [
            {"tipo":"Força","descricao":"Alta capacidade técnica","prioridade":"Alta"}
        ],
        "KPIs": [
            {
                "nome":"Aumentar faturamento",
                "area":"Comercial",
                "unidade":"R$",
                "descricao":"Aumentar vendas recorrentes em serviços de alto valor",
                "inicio_ano": datetime.now().year,
                "inicio_mes": datetime.now().month,
                "meses": [
                    {"ano": datetime.now().year + (0 if m + datetime.now().month <= 12 else 1),
                     "mes": ((datetime.now().month - 1 + m) % 12) + 1,
                     "previsto": 10000 + 500 * m,
                     "realizado": 9000 + 450 * m if m % 3 != 0 else 0.0}
                    for m in range(36)
                ]
            }
        ],
        "actions": [
            {"titulo":"Padronizar propostas","area":"Comercial","responsavel":"João","descricao":"Padronizar template e SLA",
             "data_inicio": (datetime.now()).strftime("%Y-%m-%d"), "data_vencimento": (datetime.now()).strftime("%Y-%m-%d"), "status":"Pendente","observacoes":""}
        ]
    }
    return example

def KPI_to_df(KPI: Dict[str, Any]) -> pd.DataFrame:
    rows = []
    meses = KPI.get("meses", [])
    for i, m in enumerate(meses, start=1):
        rows.append({
            "m_index": i,
            "ano": m.get("ano"),
            "mes": m.get("mes"),
            "previsto": float(m.get("previsto", 0.0)),
            "realizado": float(m.get("realizado", 0.0))
        })
    return pd.DataFrame(rows)

def planning_to_dfs(planning: Dict[str, Any]) -> Dict[str, pd.DataFrame]:
    partners = pd.DataFrame(planning.get("partners", []))
    areas = pd.DataFrame(planning.get("areas", []))
    swot = pd.DataFrame(planning.get("swot", []))
    actions = pd.DataFrame(planning.get("actions", []))
    KPI_rows = []
    KPI_mes_rows = []
    for idx, o in enumerate(planning.get("KPIs", []), start=1):
        KPI_rows.append({
            "KPI_id": idx,
            "nome": o.get("nome"),
            "area": o.get("area"),
            "unidade": o.get("unidade"),
            "descricao": o.get("descricao"),
            "inicio_ano": o.get("inicio_ano"),
            "inicio_mes": o.get("inicio_mes")
        })
        for i, m in enumerate(o.get("meses", []), start=1):
            KPI_mes_rows.append({
                "KPI_id": idx,
                "idx_mes": i,
                "ano": m.get("ano"),
                "mes": m.get("mes"),
                "previsto": float(m.get("previsto", 0.0)),
                "realizado": float(m.get("realizado", 0.0))
            })
    KPIs = pd.DataFrame(KPI_rows) if KPI_rows else pd.DataFrame(columns=["KPI_id","nome","area","unidade","descricao","inicio_ano","inicio_mes"])
    KPI_mes = pd.DataFrame(KPI_mes_rows) if KPI_mes_rows else pd.DataFrame(columns=["KPI_id","idx_mes","ano","mes","previsto","realizado"])
    return {
        "partners": partners,
        "areas": areas,
        "swot": swot,
        "KPIs": KPIs,
        "KPI_mes": KPI_mes,
        "actions": actions
    }

# -------------------------
# Gráficos (plotly -> PNG bytes)
# -------------------------

def fig_KPIs_aggregated(KPIs_df: pd.DataFrame, KPI_mes_df: pd.DataFrame) -> bytes:
    labels = []
    totals_prev = []
    totals_real = []
    for _, row in KPIs_df.iterrows():
        KPI_id = row["KPI_id"]
        dfm = KPI_mes_df[KPI_mes_df["KPI_id"] == KPI_id]
        tp = float(dfm["previsto"].sum())
        tr = float(dfm["realizado"].sum())
        labels.append(row["nome"])
        totals_prev.append(tp)
        totals_real.append(tr)
    pct_real = [(tr / tp * 100) if tp != 0 else 0.0 for tp, tr in zip(totals_prev, totals_real)]
    fig = go.Figure()
    fig.add_trace(go.Bar(x=labels, y=totals_prev, name="Total Previsto", marker_color="#4c8cff"))
    fig.add_trace(go.Bar(x=labels, y=totals_real, name="Total Realizado", marker_color="#42b983"))
    fig.add_trace(go.Scatter(x=labels, y=pct_real, name="% Realização", yaxis="y2", mode="lines+markers", line=dict(color="black", dash="dash")))
    fig.update_layout(barmode="group", yaxis=dict(title="Valor"), yaxis2=dict(title="% Realização", overlaying="y", side="right"), template="plotly_white", height=360)
    img = fig.to_image(format="png", width=1100, height=360)
    return img

def fig_KPI_monthly(KPI_df: pd.DataFrame) -> bytes:
    # KPI_df must have columns: m_index, previsto, realizado
    fig = go.Figure()
    fig.add_trace(go.Bar(x=KPI_df["m_index"], y=KPI_df["previsto"], name="Previsto", marker_color="#4c8cff"))
    fig.add_trace(go.Bar(x=KPI_df["m_index"], y=KPI_df["realizado"], name="Realizado", marker_color="#42b983"))
    y = KPI_df["realizado"].values
    x = KPI_df["m_index"].values
    if np.count_nonzero(y) >= 3:
        z = np.polyfit(x, y, 1)
        p = np.poly1d(z)
        trend = p(x)
        fig.add_trace(go.Scatter(x=x, y=trend, mode="lines", name="Tendência (Realizado)", line=dict(color="black", dash="dash")))
    fig.update_layout(barmode="group", xaxis_title="Mês (1-36)", yaxis_title="Valor", template="plotly_white", height=340)
    img = fig.to_image(format="png", width=1100, height=340)
    return img

def fig_actions_monthly(actions_df: pd.DataFrame) -> bytes:
    if actions_df.empty:
        # generate blank small image
        fig = go.Figure()
        fig.update_layout(template="plotly_white", height=300)
        return fig.to_image(format="png", width=900, height=300)
    df = actions_df.copy()
    df["data_dt"] = pd.to_datetime(df["data_vencimento"], errors="coerce")
    df["ym"] = df["data_dt"].dt.to_period("M")
    df = df.dropna(subset=["ym"])
    if df.empty:
        fig = go.Figure()
        fig.update_layout(template="plotly_white", height=300)
        return fig.to_image(format="png", width=900, height=300)
    periods = pd.period_range(start=df["ym"].min(), end=df["ym"].max(), freq="M")
    labels = [p.strftime("%Y-%m") for p in periods]
    total_due = []
    pct_done = []
    for p in periods:
        sel = df[df["ym"] == p]
        total = len(sel)
        done = len(sel[sel["status"] == "Concluído"])
        total_due.append(total)
        pct_done.append((done / total * 100) if total > 0 else 0.0)
    fig = go.Figure()
    fig.add_trace(go.Bar(x=labels, y=total_due, name="Planos com vencimento", marker_color="#ff7f0e"))
    fig.add_trace(go.Scatter(x=labels, y=pct_done, name="% concluídos", yaxis="y2", mode="lines+markers", line=dict(color="black", dash="dash")))
    fig.update_layout(xaxis_tickangle=-45, yaxis=dict(title="Qtde"), yaxis2=dict(overlaying="y", side="right", title="% concluídos"), template="plotly_white", height=340)
    return fig.to_image(format="png", width=1100, height=340)

# -------------------------
# Word generation
# -------------------------

def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(14 if level == 1 else 12)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_paragraph(doc: Document, text: str, bold: bool = False, size: int = 11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Segoe UI"
    r.font.size = Pt(size)
    return p

def add_table_from_df(doc: Document, df: pd.DataFrame, title: str = None):
    if title:
        add_paragraph(doc, title, bold=True, size=12)
    if df.empty:
        add_paragraph(doc, "(nenhum registro)", size=10)
        return
    # Create table with header
    table = doc.add_table(rows=1, cols=len(df.columns), style="Table Grid")
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = "Segoe UI"
                r.font.size = Pt(10)
                r.bold = True
    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if pd.isna(val):
                txt = ""
            else:
                if isinstance(val, float):
                    txt = f"{val:,.2f}"
                else:
                    txt = str(val)
            row_cells[i].text = txt
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Segoe UI"
                    r.font.size = Pt(10)
    doc.add_paragraph("")  # spacing

def insert_image_bytes(doc: Document, image_bytes: bytes, width_inches: float = 6.5):
    bio = io.BytesIO(image_bytes)
    doc.add_picture(bio, width=Inches(width_inches))

def generate_docx(planning_data: Dict[str, Any], output_path: str = "relatorio_planejamento.docx"):
    dfs = planning_to_dfs(planning_data)
    doc = Document()
    # Title
    add_heading(doc, "BK_Planejamento_Estrategico - Relatório de Planejamento Estratégico", level=1)
    add_paragraph(doc, f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", size=10)

    # Partners
    add_heading(doc, "Sócios", level=2)
    add_table_from_df(doc, dfs["partners"], title=None)

    # Areas
    add_heading(doc, "Áreas e Responsáveis", level=2)
    add_table_from_df(doc, dfs["areas"])

    # SWOT
    add_heading(doc, "SWOT", level=2)
    add_table_from_df(doc, dfs["swot"])

    # KPIs aggregated
    add_heading(doc, "KPIs (3 anos) - Visão Geral", level=2)
    if not dfs["KPIs"].empty:
        try:
            agg_img = fig_KPIs_aggregated(dfs["KPIs"], dfs["KPI_mes"])
            insert_image_bytes(doc, agg_img, width_inches=6.5)
        except Exception as e:
            add_paragraph(doc, f"Erro ao gerar gráfico agregado: {e}", size=10)
    else:
        add_paragraph(doc, "Nenhuma KPI cadastrada.", size=10)

    # Per-KPI details
    for _, KPI_row in dfs["KPIs"].iterrows():
        KPI_id = KPI_row["KPI_id"]
        add_heading(doc, f"KPI: {KPI_row['nome']}", level=3)
        add_paragraph(doc, f"Área: {KPI_row.get('area','')}. Unidade: {KPI_row.get('unidade','')}. Início: {int(KPI_row.get('inicio_mes',1)):02d}/{KPI_row.get('inicio_ano','')}", size=10)
        if KPI_row.get("descricao"):
            add_paragraph(doc, KPI_row["descricao"], size=10)
        dfm = dfs["KPI_mes"][dfs["KPI_mes"]["KPI_id"] == KPI_id].copy().sort_values("idx_mes")
        if dfm.empty:
            add_paragraph(doc, "(nenhum mês cadastrado)", size=10)
            continue
        # Table of months
        dfm_display = dfm[["idx_mes","ano","mes","previsto","realizado"]].rename(columns={"idx_mes":"Mês idx","ano":"Ano","mes":"Mês","previsto":"Previsto","realizado":"Realizado"})
        add_table_from_df(doc, dfm_display)
        # chart
        try:
            img = fig_KPI_monthly(dfm)
            insert_image_bytes(doc, img, width_inches=6.5)
        except Exception as e:
            add_paragraph(doc, f"Erro ao gerar gráfico mensal: {e}", size=10)
        # Recommendations (simple)
        try:
            avg_prev = dfm["previsto"].mean() if not dfm["previsto"].isna().all() else 0
            avg_real = dfm["realizado"].mean() if not dfm["realizado"].isna().all() else 0
            recs = []
            if avg_prev == 0:
                recs.append("Definir metas previstas (previsto=0 impede análise).")
            else:
                pct = ((avg_real - avg_prev) / avg_prev) * 100
                if pct < -10:
                    recs.append("Realizado consistentemente abaixo do previsto (>10%). Revisar causas.")
                elif pct < 0:
                    recs.append("Leve subperformance. Reforçar acompanhamento.")
                elif pct < 10:
                    recs.append("Performance adequada. Padronizar processos.")
                else:
                    recs.append("Realizado acima do previsto — validar sustentabilidade e ajustar metas se necessário.")
                recs.append("Alinhar KPI com planos de ação e responsáveis com datas claras.")
            add_paragraph(doc, "Recomendações:", bold=True)
            for r in recs:
                add_paragraph(doc, f"- {r}", size=10)
        except Exception:
            pass

    # Planos de Ação
    add_heading(doc, "Planos de Ação", level=2)
    add_table_from_df(doc, dfs["actions"])
    try:
        img_actions = fig_actions_monthly(dfs["actions"])
        insert_image_bytes(doc, img_actions, width_inches=6.5)
    except Exception as e:
        add_paragraph(doc, f"Erro ao gerar gráfico de planos: {e}", size=10)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Produzido por BK Engenharia e Tecnologia")
    r.font.name = "Segoe UI"
    r.font.size = Pt(9)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save
    doc.save(output_path)
    print(f"Relatório salvo em: {output_path}")

# -------------------------
# Execução
# -------------------------

if __name__ == "__main__":
    planning = load_planning_json("planning.json")
    generate_docx(planning, "relatorio_planejamento.docx")
